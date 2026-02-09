#!/usr/bin/env python3
"""
Generate DOCX templates for French administrative documents.
Templates use {{placeholder}} syntax for variable substitution.
Based on official government structures (legifrance.gouv.fr, collectivites-locales.gouv.fr).
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'templates_docx')
os.makedirs(OUTPUT_DIR, exist_ok=True)


def setup_styles(doc):
    """Configure French administrative document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_header(doc, title):
    """Add Republic header."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('RÉPUBLIQUE FRANÇAISE')
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Liberté - Égalité - Fraternité')
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('COMMUNE DE {{commune_nom}}')
    run.bold = True
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Département de {{departement_nom}}')
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.underline = True


def add_signature(doc):
    """Add standard signature block."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run('Fait à {{commune_nom}}, le {{date_signature}}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Le Maire,')
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('{{maire_nom}}')


def create_arrete_voirie():
    """Create Arrêté de voirie template."""
    doc = Document()
    setup_styles(doc)
    add_header(doc, 'ARRÊTÉ DE VOIRIE N° {{numero_arrete}}')

    # VU section
    doc.add_paragraph()
    visas = [
        "VU le Code général des collectivités territoriales, notamment les articles L.2213-1 à L.2213-6 ;",
        "VU le Code de la route, notamment les articles R.411-5, R.411-8, R.411-18, R.411-25 à R.411-28 ;",
        "VU l'arrêté interministériel du 24 novembre 1967 relatif à la signalisation des routes et des autoroutes ;",
        "VU l'instruction interministérielle sur la signalisation routière ;",
        "VU le pouvoir de police du maire en matière de circulation et de stationnement ;",
    ]
    for v in visas:
        p = doc.add_paragraph(v)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # CONSIDERANT section
    considerants = [
        "CONSIDÉRANT que des travaux de {{nature_travaux}} doivent être réalisés sur la voie {{rue_nom}}, entre le n° {{numero_debut}} et le n° {{numero_fin}} ;",
        "CONSIDÉRANT qu'il convient, pour la sécurité des usagers de la voie publique et des personnels de chantier, de réglementer temporairement la circulation et le stationnement ;",
        "CONSIDÉRANT la demande présentée par {{demandeur_nom}} en date du {{date_demande}} ;",
    ]
    for c in considerants:
        p = doc.add_paragraph(c)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ARRETE section
    p = doc.add_paragraph()
    run = p.add_run('ARRÊTE :')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    articles = [
        ("Article 1er – Objet",
         "À compter du {{date_debut}} jusqu'au {{date_fin}} inclus, la circulation sera {{type_restriction}} sur la voie {{rue_nom}}, entre le n° {{numero_debut}} et le n° {{numero_fin}}, pour permettre l'exécution de travaux de {{nature_travaux}}."),
        ("Article 2 – Déviation",
         "Une déviation sera mise en place via {{itineraire_deviation}}. La signalisation de déviation sera conforme aux dispositions réglementaires en vigueur."),
        ("Article 3 – Stationnement",
         "Le stationnement est interdit sur la zone de travaux et sur {{zone_stationnement_interdit}} pendant toute la durée des travaux, sauf pour les véhicules de chantier dûment autorisés."),
        ("Article 4 – Vitesse",
         "La vitesse est limitée à {{vitesse_limitee}} km/h dans la zone de chantier."),
        ("Article 5 – Signalisation",
         "L'entreprise {{entreprise_nom}} est chargée de la fourniture, de la mise en place et de l'entretien de la signalisation temporaire conformément à l'instruction interministérielle sur la signalisation routière."),
        ("Article 6 – Publication",
         "Le présent arrêté sera affiché aux extrémités de la zone réglementée et publié au recueil des actes administratifs de la commune."),
        ("Article 7 – Recours",
         "Le présent arrêté peut faire l'objet d'un recours contentieux devant le tribunal administratif de {{tribunal_administratif}} dans un délai de deux mois à compter de sa publication ou de sa notification."),
        ("Article 8 – Exécution",
         "Le secrétaire général de la mairie, le commandant de la brigade de gendarmerie et tout agent de la force publique sont chargés, chacun en ce qui le concerne, de l'exécution du présent arrêté."),
    ]

    for title, text in articles:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.underline = True
        doc.add_paragraph(text)

    add_signature(doc)

    path = os.path.join(OUTPUT_DIR, 'arrete_voirie.docx')
    doc.save(path)
    print(f"Created: {path}")
    return path


def create_arrete_police():
    """Create Arrêté de police du maire template."""
    doc = Document()
    setup_styles(doc)
    add_header(doc, 'ARRÊTÉ DE POLICE N° {{numero_arrete}}')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('{{objet_arrete}}')
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    # VU section
    visas = [
        "VU le Code général des collectivités territoriales, notamment les articles L.2212-1 et L.2212-2 ;",
        "VU le Code de la santé publique ;",
        "VU le Code de l'environnement ;",
        "VU le Code pénal, notamment les articles R.610-5 et R.644-2 ;",
    ]
    for v in visas:
        p = doc.add_paragraph(v)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # CONSIDERANT section
    considerants = [
        "CONSIDÉRANT qu'il appartient au maire d'assurer le bon ordre, la sûreté, la sécurité et la salubrité publiques sur le territoire de la commune ;",
        "CONSIDÉRANT que {{motif_principal}} ;",
        "CONSIDÉRANT la nécessité de prévenir {{risque_ou_nuisance}} ;",
        "CONSIDÉRANT que {{justification_complementaire}} ;",
    ]
    for c in considerants:
        p = doc.add_paragraph(c)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('ARRÊTE :')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    articles = [
        ("Article 1er – Mesures prescrites",
         "{{mesures_prescrites}}"),
        ("Article 2 – Périmètre d'application",
         "Les mesures prévues à l'article 1er s'appliquent sur le périmètre suivant : {{perimetre_application}}."),
        ("Article 3 – Durée d'application",
         "Le présent arrêté entre en vigueur le {{date_debut}} et reste applicable jusqu'au {{date_fin}}."),
        ("Article 4 – Sanctions",
         "Toute infraction au présent arrêté sera constatée et poursuivie conformément aux lois et règlements en vigueur. Les contrevenants s'exposent aux sanctions prévues par l'article R.610-5 du Code pénal."),
        ("Article 5 – Publication et notification",
         "Le présent arrêté sera affiché en mairie, publié au recueil des actes administratifs de la commune et notifié aux personnes intéressées."),
        ("Article 6 – Recours",
         "Le présent arrêté peut faire l'objet d'un recours contentieux devant le tribunal administratif de {{tribunal_administratif}} dans un délai de deux mois à compter de sa publication ou de sa notification."),
        ("Article 7 – Exécution",
         "Le secrétaire général de la mairie, les agents de police municipale et tout agent de la force publique sont chargés, chacun en ce qui le concerne, de l'exécution du présent arrêté."),
    ]

    for title, text in articles:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.underline = True
        doc.add_paragraph(text)

    add_signature(doc)

    path = os.path.join(OUTPUT_DIR, 'arrete_police_maire.docx')
    doc.save(path)
    print(f"Created: {path}")
    return path


def create_permis_stationnement():
    """Create Permis de stationnement template."""
    doc = Document()
    setup_styles(doc)
    add_header(doc, 'ARRÊTÉ DE VOIRIE N° {{numero_arrete}}\nPORTANT PERMIS DE STATIONNEMENT')

    doc.add_paragraph()

    # VU section
    visas = [
        "VU le Code général des collectivités territoriales, notamment les articles L.2213-1 et L.2213-6 ;",
        "VU le Code de la voirie routière ;",
        "VU le Code général de la propriété des personnes publiques, notamment les articles L.2122-1 et suivants ;",
        "VU le règlement de voirie communale ;",
        "VU la demande présentée par {{demandeur_nom}} en date du {{date_demande}} ;",
    ]
    for v in visas:
        p = doc.add_paragraph(v)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    considerants = [
        "CONSIDÉRANT que l'occupation du domaine public sollicitée est compatible avec la conservation du domaine public et l'exercice des activités de service public ;",
        "CONSIDÉRANT que l'installation de {{objet_occupation}} ne constitue pas un obstacle à la circulation des piétons et des véhicules ;",
    ]
    for c in considerants:
        p = doc.add_paragraph(c)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('ARRÊTE :')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    articles = [
        ("Article 1er – Bénéficiaire",
         "{{demandeur_nom}}, demeurant {{demandeur_adresse}}, est autorisé(e) à occuper temporairement le domaine public communal dans les conditions définies ci-après."),
        ("Article 2 – Objet et localisation",
         "L'autorisation porte sur l'installation de {{objet_occupation}} au droit du {{adresse_occupation}}, sur une superficie de {{surface_m2}} m²."),
        ("Article 3 – Durée",
         "La présente autorisation est accordée du {{date_debut}} au {{date_fin}}."),
        ("Article 4 – Caractère précaire et révocable",
         "La présente autorisation est délivrée à titre précaire et révocable. Elle peut être retirée à tout moment, pour un motif d'intérêt général, sans que le bénéficiaire puisse prétendre à une quelconque indemnité."),
        ("Article 5 – Redevance",
         "L'occupation du domaine public donne lieu au paiement d'une redevance d'un montant de {{montant_redevance}} euros, conformément au tarif en vigueur adopté par le conseil municipal."),
        ("Article 6 – Prescriptions techniques",
         "Le bénéficiaire devra respecter les prescriptions suivantes :\n- Maintenir la propreté des lieux occupés et de leurs abords ;\n- Ne pas gêner la circulation des piétons et des véhicules ;\n- {{prescriptions_supplementaires}} ;\n- Respecter les horaires suivants : {{horaires_autorises}}."),
        ("Article 7 – Signalisation",
         "Le bénéficiaire est tenu de mettre en place, à ses frais, la signalisation réglementaire nécessaire."),
        ("Article 8 – Responsabilité et assurances",
         "Le bénéficiaire est responsable de tout dommage causé aux tiers ou au domaine public du fait de l'occupation autorisée. Il devra justifier d'une assurance couvrant sa responsabilité civile."),
        ("Article 9 – Remise en état",
         "À l'expiration de l'autorisation ou en cas de retrait, le bénéficiaire devra remettre les lieux dans leur état initial, dans un délai de {{delai_remise_etat}} jours."),
        ("Article 10 – Publication et notification",
         "Le présent arrêté sera affiché sur les lieux et notifié au bénéficiaire. Il sera également publié au recueil des actes administratifs."),
        ("Article 11 – Recours",
         "Le présent arrêté peut faire l'objet d'un recours contentieux devant le tribunal administratif de {{tribunal_administratif}} dans un délai de deux mois à compter de sa publication ou de sa notification."),
        ("Article 12 – Exécution",
         "Le secrétaire général de la mairie et les services municipaux compétents sont chargés, chacun en ce qui le concerne, de l'exécution du présent arrêté."),
    ]

    for title, text in articles:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.underline = True
        doc.add_paragraph(text)

    add_signature(doc)

    path = os.path.join(OUTPUT_DIR, 'permis_stationnement.docx')
    doc.save(path)
    print(f"Created: {path}")
    return path


def create_placeholder_metadata():
    """Generate JSON metadata for each template's placeholders."""
    import json

    metadata = {
        "arrete_voirie": {
            "nom": "Arrêté de voirie",
            "categorie": "arrete_municipal",
            "description": "Arrêté municipal réglementant temporairement la circulation et le stationnement pour travaux sur la voie publique",
            "references_legales": ["L.2213-1 CGCT", "L.2213-6 CGCT", "R.411-5 Code de la route"],
            "placeholders": [
                {"id": "commune_nom", "label": "Nom de la commune", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "departement_nom", "label": "Nom du département", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "numero_arrete", "label": "Numéro de l'arrêté", "type": "text", "required": True, "example": "2026-ARR-042"},
                {"id": "nature_travaux", "label": "Nature des travaux", "type": "text", "required": True, "example": "réfection de la chaussée", "question_chatbot": "Quelle est la nature des travaux à réaliser ?"},
                {"id": "rue_nom", "label": "Nom de la rue concernée", "type": "text", "required": True, "question_chatbot": "Sur quelle rue les travaux auront-ils lieu ?"},
                {"id": "numero_debut", "label": "Numéro de début de la zone", "type": "text", "required": True, "question_chatbot": "À quel numéro commence la zone de travaux ?"},
                {"id": "numero_fin", "label": "Numéro de fin de la zone", "type": "text", "required": True, "question_chatbot": "À quel numéro se termine la zone de travaux ?"},
                {"id": "demandeur_nom", "label": "Nom du demandeur", "type": "text", "required": True, "question_chatbot": "Qui est le demandeur des travaux (entreprise ou particulier) ?"},
                {"id": "date_demande", "label": "Date de la demande", "type": "date", "required": True, "question_chatbot": "Quelle est la date de la demande ?"},
                {"id": "date_debut", "label": "Date de début des travaux", "type": "date", "required": True, "question_chatbot": "Quelle est la date de début des travaux ?"},
                {"id": "date_fin", "label": "Date de fin des travaux", "type": "date", "required": True, "question_chatbot": "Quelle est la date prévue de fin des travaux ?"},
                {"id": "type_restriction", "label": "Type de restriction de circulation", "type": "enum", "required": True, "options": ["interrompue", "alternée", "réduite à une voie", "déviée"], "question_chatbot": "Quel type de restriction de circulation souhaitez-vous mettre en place ?"},
                {"id": "itineraire_deviation", "label": "Itinéraire de déviation", "type": "text", "required": False, "question_chatbot": "Quel est l'itinéraire de déviation prévu (si applicable) ?"},
                {"id": "zone_stationnement_interdit", "label": "Zone de stationnement interdit", "type": "text", "required": False, "default": "les abords immédiats du chantier"},
                {"id": "vitesse_limitee", "label": "Vitesse limitée (km/h)", "type": "number", "required": False, "default": "30"},
                {"id": "entreprise_nom", "label": "Nom de l'entreprise chargée de la signalisation", "type": "text", "required": True, "question_chatbot": "Quelle entreprise sera chargée de la signalisation ?"},
                {"id": "tribunal_administratif", "label": "Tribunal administratif compétent", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "date_signature", "label": "Date de signature", "type": "date", "required": True, "auto_fill": True, "source": "today"},
                {"id": "maire_nom", "label": "Nom du maire", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
            ]
        },
        "arrete_police_maire": {
            "nom": "Arrêté de police du maire",
            "categorie": "arrete_municipal",
            "description": "Arrêté municipal pris au titre des pouvoirs de police générale du maire pour assurer l'ordre, la sûreté, la sécurité et la salubrité publiques",
            "references_legales": ["L.2212-1 CGCT", "L.2212-2 CGCT", "R.610-5 Code pénal"],
            "placeholders": [
                {"id": "commune_nom", "label": "Nom de la commune", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "departement_nom", "label": "Nom du département", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "numero_arrete", "label": "Numéro de l'arrêté", "type": "text", "required": True, "example": "2026-POL-015"},
                {"id": "objet_arrete", "label": "Objet de l'arrêté", "type": "text", "required": True, "question_chatbot": "Quel est l'objet de cet arrêté de police ? (ex: réglementation du bruit, interdiction de circulation, etc.)"},
                {"id": "motif_principal", "label": "Motif principal", "type": "text", "required": True, "question_chatbot": "Quel est le motif principal justifiant cet arrêté ?"},
                {"id": "risque_ou_nuisance", "label": "Risque ou nuisance à prévenir", "type": "text", "required": True, "question_chatbot": "Quel risque ou quelle nuisance cet arrêté vise-t-il à prévenir ?"},
                {"id": "justification_complementaire", "label": "Justification complémentaire", "type": "text", "required": False, "question_chatbot": "Y a-t-il une justification complémentaire à ajouter ?"},
                {"id": "mesures_prescrites", "label": "Mesures prescrites", "type": "text", "required": True, "question_chatbot": "Quelles sont les mesures prescrites par cet arrêté ? (décrivez précisément les interdictions ou obligations)"},
                {"id": "perimetre_application", "label": "Périmètre d'application", "type": "text", "required": True, "question_chatbot": "Quel est le périmètre géographique d'application ?"},
                {"id": "date_debut", "label": "Date d'entrée en vigueur", "type": "date", "required": True, "question_chatbot": "À quelle date l'arrêté entre-t-il en vigueur ?"},
                {"id": "date_fin", "label": "Date de fin de validité", "type": "date", "required": False, "question_chatbot": "Quelle est la date de fin de validité ? (laisser vide si permanent)"},
                {"id": "tribunal_administratif", "label": "Tribunal administratif compétent", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "date_signature", "label": "Date de signature", "type": "date", "required": True, "auto_fill": True, "source": "today"},
                {"id": "maire_nom", "label": "Nom du maire", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
            ]
        },
        "permis_stationnement": {
            "nom": "Permis de stationnement",
            "categorie": "autorisation",
            "description": "Autorisation d'occupation temporaire du domaine public communal (terrasse, échafaudage, étalage, etc.)",
            "references_legales": ["L.2213-1 CGCT", "L.2213-6 CGCT", "L.2122-1 CGPPP"],
            "placeholders": [
                {"id": "commune_nom", "label": "Nom de la commune", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "departement_nom", "label": "Nom du département", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "numero_arrete", "label": "Numéro de l'arrêté", "type": "text", "required": True},
                {"id": "demandeur_nom", "label": "Nom du demandeur", "type": "text", "required": True, "question_chatbot": "Quel est le nom du demandeur (personne ou entreprise) ?"},
                {"id": "demandeur_adresse", "label": "Adresse du demandeur", "type": "text", "required": True, "question_chatbot": "Quelle est l'adresse complète du demandeur ?"},
                {"id": "date_demande", "label": "Date de la demande", "type": "date", "required": True, "question_chatbot": "Quelle est la date de la demande d'occupation ?"},
                {"id": "objet_occupation", "label": "Objet de l'occupation", "type": "text", "required": True, "question_chatbot": "Quel est l'objet de l'occupation du domaine public ? (terrasse, échafaudage, étalage, etc.)"},
                {"id": "adresse_occupation", "label": "Adresse de l'occupation", "type": "text", "required": True, "question_chatbot": "À quelle adresse se situe l'emplacement demandé ?"},
                {"id": "surface_m2", "label": "Surface en m²", "type": "number", "required": True, "question_chatbot": "Quelle est la superficie occupée en m² ?"},
                {"id": "date_debut", "label": "Date de début d'occupation", "type": "date", "required": True, "question_chatbot": "À quelle date commence l'occupation ?"},
                {"id": "date_fin", "label": "Date de fin d'occupation", "type": "date", "required": True, "question_chatbot": "À quelle date se termine l'occupation ?"},
                {"id": "montant_redevance", "label": "Montant de la redevance (€)", "type": "number", "required": True, "question_chatbot": "Quel est le montant de la redevance d'occupation ?"},
                {"id": "prescriptions_supplementaires", "label": "Prescriptions supplémentaires", "type": "text", "required": False, "question_chatbot": "Y a-t-il des prescriptions techniques supplémentaires ?"},
                {"id": "horaires_autorises", "label": "Horaires autorisés", "type": "text", "required": False, "default": "tous les jours de 8h00 à 22h00"},
                {"id": "delai_remise_etat", "label": "Délai de remise en état (jours)", "type": "number", "required": False, "default": "15"},
                {"id": "tribunal_administratif", "label": "Tribunal administratif compétent", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
                {"id": "date_signature", "label": "Date de signature", "type": "date", "required": True, "auto_fill": True, "source": "today"},
                {"id": "maire_nom", "label": "Nom du maire", "type": "text", "required": True, "auto_fill": True, "source": "commune"},
            ]
        }
    }

    path = os.path.join(OUTPUT_DIR, 'templates_metadata.json')
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)
    print(f"Created: {path}")
    return metadata


if __name__ == '__main__':
    create_arrete_voirie()
    create_arrete_police()
    create_permis_stationnement()
    create_placeholder_metadata()
    print("\nAll templates created successfully!")
